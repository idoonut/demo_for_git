from configparser import SectionProxy
from Common.Backend import Get_Billsday, Get_Sale_store, Get_opn
from datetime import date,timedelta
from docx import Document
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

dte=date.today()

WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"


def set_number_of_columns(section, cols):
    """ sets number of columns through xpath. """
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(cols))

def del_para(para):
    p=para._element
    p.getparent().remove(p)
    p._p=p._element=None

def gen_billtoday(today):
    from docx2pdf import convert
    import os
    db=Get_Billsday(today)
    opn=Get_opn(today)
    print(db)
    billdoc=Document()
    sec1=billdoc.add_section(WD_SECTION.ODD_PAGE)
    set_number_of_columns(sec1, 2)
    opening=billdoc.add_heading(f'Opening Cash-{opn}', level=2)
    SSh=billdoc.add_heading('Silver Sale', level=2)
    SSp=billdoc.add_table(rows=1, cols=4)
    SSp.style='TableGrid'
    SSpr=SSp.rows[0].cells
    SSpr[0].text='Amount'
    SSpr[1].text='Item'
    SSpr[2].text='Gr Wt'
    SSpr[3].text='Nt Wt'
    SHPh=billdoc.add_heading('Misc/Shree Aaya', level=2)
    SHPp=billdoc.add_table(rows=1, cols=2)
    SHPp.style='TableGrid'
    SHPp.rows[0].cells[0].text='Amount'
    SHPp.rows[0].cells[1].text='Type'
    OOh=billdoc.add_heading('Old Order', level=2)
    OOp=billdoc.add_table(rows=1, cols=5)
    OOp.style='TableGrid'
    OOp.rows[0].cells[0].text='Amount'
    OOp.rows[0].cells[1].text='Cust Name'
    OOp.rows[0].cells[2].text='Item'
    OOp.rows[0].cells[3].text='Gr Wt'
    OOp.rows[0].cells[4].text='Nt Wt'
    GSh=billdoc.add_heading('Gold Sale', level=2)
    GSp=billdoc.add_table(rows=1, cols=5)
    GSp.style='TableGrid'
    GSp.rows[0].cells[0].text='Item'
    GSp.rows[0].cells[1].text='Gr Wt'
    GSp.rows[0].cells[2].text='Nt Wt'
    GSp.rows[0].cells[3].text='Name'
    GSp.rows[0].cells[4].text='Amount'
    NOh=billdoc.add_heading('New Order', level=2)
    NOp=billdoc.add_table(rows=1, cols=6)
    NOp.style='TableGrid'
    NOp.rows[0].cells[0].text='Est Amt'
    NOp.rows[0].cells[1].text='Adv'
    NOp.rows[0].cells[2].text='Name'
    NOp.rows[0].cells[3].text='Item'
    NOp.rows[0].cells[4].text='Gr Wt'
    NOp.rows[0].cells[5].text='Nt Wt'
    VCh=billdoc.add_heading('VC', level=2)
    VCp=billdoc.add_table(rows=1, cols=5)
    VCp.style='TableGrid'
    VCp.rows[0].cells[0].text='Amount'
    VCp.rows[0].cells[1].text='Name'
    VCp.rows[0].cells[2].text='Item'
    VCp.rows[0].cells[3].text='Gr Wt'
    VCp.rows[0].cells[4].text='Nt Wt'
    section=billdoc.add_section(WD_SECTION.NEW_COLUMN)
    SPh=billdoc.add_heading('Silver Purchase', level=2)
    SPp=billdoc.add_table(rows=1, cols=4)
    SPp.style='TableGrid'
    SPp.rows[0].cells[0].text='Amount'
    SPp.rows[0].cells[1].text='Item'
    SPp.rows[0].cells[2].text='Gr Wt'
    SPp.rows[0].cells[3].text='Nt Wt'
    SHMh=billdoc.add_heading('Misc/Shree Gaya', level=2)
    SHMp=billdoc.add_table(rows=1, cols=2)
    SHMp.style='TableGrid'
    SHMp.rows[0].cells[0].text='Amount'
    SHMp.rows[0].cells[1].text='Type'
    BEh=billdoc.add_heading('Bank Entries', level=2)
    BEp=billdoc.add_table(rows=1, cols=3)
    BEp.style='TableGrid'
    BEp.rows[0].cells[0].text='Amount'
    BEp.rows[0].cells[1].text='Type'
    BEp.rows[0].cells[2].text='Name'
    Exh=billdoc.add_heading('Expenses', level=2)
    Exp=billdoc.add_table(rows=1, cols=2)
    Exp.style='TableGrid'
    Exp.rows[0].cells[0].text='Amount'
    Exp.rows[0].cells[1].text='Type'
    GPh=billdoc.add_heading('Gold Purchase', level=2)
    GPp=billdoc.add_table(rows=1, cols=5)
    GPp.style='TableGrid'
    GPp.rows[0].cells[0].text='Item'
    GPp.rows[0].cells[1].text='Gr Wt'
    GPp.rows[0].cells[2].text='Nt Wt'
    GPp.rows[0].cells[3].text='Name'
    GPp.rows[0].cells[4].text='Amount'
    CCh=billdoc.add_heading('Closing Cash', level=2)
    GStotal,SStotal, OOtotal, NOtotal, VCtotal, GPtotal, SPtotal, SHPtotal, SHMtotal, EXtotal=0,0,0,0,0,0,0,0,0,0
    SSntwt, GSntwt, GPntwt, SPntwt=0.0, 0.0, 0.0, 0.0
    for i in db:
        print(i)
        if i[10]=='Gold Sale':
            GSrow=GSp.add_row().cells
            GSrow[0].text=str(i[3])
            GSrow[1].text=str(i[6])
            GSrow[2].text=str(i[7])
            GSrow[3].text=str(i[0])
            GSrow[4].text=str(i[11])
            GStotal+=i[11]
            GSntwt+=i[7]
        elif i[10]=='Silver Sale':
            SSrow=SSp.add_row().cells
            SSrow[0].text=str(i[11])
            SSrow[1].text=str(i[3])
            SSrow[2].text=str(i[8])
            SSrow[3].text=str(i[9])
            SStotal+=i[11]
            SSntwt+=i[7]
        elif i[10]=='Old Order':
            OOrow=OOp.add_row().cells
            OOrow[0].text=str(i[11])
            OOrow[1].text=str(i[0])
            OOrow[2].text=str(i[3])
            OOrow[3].text=str(i[6])
            OOrow[4].text=str(i[7])
            OOtotal+=i[11]
        elif i[10]=='New Order':
            NOrow=NOp.add_row().cells
            NOrow[0].text=str(i[4])
            NOrow[1].text=str(i[11])
            NOrow[2].text=str(i[0])
            NOrow[3].text=str(i[3])
            NOrow[4].text=str(i[6])
            NOrow[5].text=str(i[7])
            NOtotal+=i[11]
        elif i[10]=='VC':
            VCrow=VCp.add_row().cells
            VCrow[0].text=str(i[11])
            VCrow[1].text=str(i[0])
            VCrow[2].text=str(i[3])
            VCrow[3].text=str(i[6])
            VCrow[4].text=str(i[7])
            VCtotal+=i[11]
        elif i[10]=='Gold Purchase':
            GProw=GPp.add_row().cells
            GProw[0].text=str(i[3])
            GProw[1].text=str(i[6])
            GProw[2].text=str(i[7])
            GProw[3].text=str(i[0])
            GProw[4].text=str(i[11])
            GPtotal+=i[11]
            GPntwt+=i[7]
        elif i[10]=='Silver Purchase':
            SProw=SPp.add_row().cells
            SProw[0].text=str(i[11])
            SProw[1].text=str(i[3])
            SProw[2].text=str(i[8])
            SProw[3].text=str(i[9])
            SPtotal+=i[11]
            SPntwt+=i[7]
        elif i[10]=='Misc Plus':
            SHProw=SHPp.add_row().cells
            SHProw[0].text=str(i[11])
            SHProw[1].text=str(i[0])
            SHPtotal+=i[11]
        elif i[10]=='Misc Minus':
            SHMrow=SHMp.add_row().cells
            SHMrow[0].text=str(i[11])
            SHMrow[1].text=str(i[0])
            SHMtotal+=i[11]
        elif i[10]=='Expense':
            Exprow=Exp.add_row().cells
            Exprow[0].text=str(i[11])
            Exprow[1].text=str(i[0])
            EXtotal+=i[11]
        if i[19]>0:
            SHMrow=SHMp.add_row().cells
            SHMrow[0].text=str(i[11])
            SHMrow[1].text=str(i[0]+" (CR)")
            SHMtotal+=i[11]
        if i[13]!=0:
            BEprow=BEp.add_row().cells
            BEprow[0].text=str(i[13])
            BEprow[1].text=str(i[14])
            BEprow[2].text=str(i[0])
        if i[15]!=0:
            BEprow=BEp.add_row().cells
            BEprow[0].text=str(i[15])
            BEprow[1].text=str(i[16])
            BEprow[2].text=str(i[0])
        if i[17]!=0:
            BEprow=BEp.add_row().cells
            BEprow[0].text=str(i[17])
            BEprow[1].text=f'Chq-{i[18]}'
            BEprow[2].text=str(i[0])


    GSh.add_run(f' - {GStotal}')
    GSlastline=GSp.add_row().cells
    GSlastline[0].text='Total'
    GSlastline[2].text=str(GSntwt)
    SSh.add_run(f' - {SStotal}')
    SSlastline=SSp.add_row().cells
    SSlastline[0].text='Total'
    SSlastline[3].text=str(SSntwt)
    OOh.add_run(f' - {OOtotal}')
    NOh.add_run(f' - {NOtotal}')
    VCh.add_run(f' - {VCtotal}')
    GPh.add_run(f' - {GPtotal}')
    GPlastline=GPp.add_row().cells
    GPlastline[0].text='Total'
    GPlastline[2].text=str(GPntwt)
    SPh.add_run(f' - {SPtotal}')
    SPlastline=SPp.add_row().cells
    SPlastline[0].text='Total'
    SPlastline[3].text=str(SPntwt)
    SHPh.add_run(f' -  {SHPtotal}')
    SHMh.add_run(f' - {SHMtotal}')
    Exh.add_run(f' - {EXtotal}')

    para1=billdoc.paragraphs[0]
    del_para(para1)
    try:
        billdoc.save(f'{today}.docx')
    except:
        import os
        os.rmdir(f'{today}.docx')
        billdoc.save(f'{today}.docx')

def gen_sum(today):
    dbdict=Get_Sale_store(today, today)
    
    sumdoc=Document()
    sumdoc.add_heading("Gold Sale", level=2)
    GS=sumdoc.add_table(rows=1, cols=2)
    GS.style='TableGrid'
    GS.rows[0].cells[0].text='Item Name'
    GS.rows[0].cells[1].text='Amount'
    sumdoc.add_heading("Silver Sale", level=2)
    SS=sumdoc.add_table(rows=1, cols=2)
    SS.style='TableGrid'
    SS.rows[0].cells[0].text='Item Name'
    SS.rows[0].cells[1].text='Amount'
    GSdup, SSdup={}, {}
    for key in dbdict.keys():
        if key=='Gold Sale':
            for i in dbdict[key]:
                for j in i:
                    if str(j) in GSdup.keys():
                        GSdup[str(j[0])]+=1
                    else:
                        GSdup[str(j[0])]=1
        if key=='Silver Sale':
            for i in dbdict[key]:
                for j in i:
                    if str(j) in SSdup.keys():
                        SSdup[str(j[0])]+=1
                    else:
                        SSdup[str(j[0])]=1
    for key in GSdup.keys():
        GSrow=GS.add_row().cells
        GSrow[0].text=key
        GSrow[1].text=str(GSdup[key])
    for key in SSdup.keys():
        SSrow=SS.add_row().cells
        SSrow[0].text=key
        SSrow[1].text=str(SSdup[key])
    sumdoc.save(f'{today}Sum.docx')

if __name__=='__main__':
    # gen_billtoday(dte)
    gen_sum(dte)